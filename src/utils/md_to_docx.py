#!/usr/bin/env python3
"""
Convert Markdown files to Word documents (.docx)

This utility converts Markdown (.md) files to Microsoft Word format (.docx)
for easier reading by non-technical users.

Usage:
    python src/utils/md_to_docx.py INSTRUCTIONS.md
    python src/utils/md_to_docx.py README.md -o output.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_to_docx(md_file, docx_file):
    """
    Convert a Markdown file to a Word document.

    Args:
        md_file: Path to input Markdown file
        docx_file: Path to output Word document
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Heading 1
        if line.startswith('# '):
            text = line[2:]
            heading = doc.add_heading(text, level=1)
            heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)

        # Heading 2
        elif line.startswith('## '):
            text = line[3:]
            heading = doc.add_heading(text, level=2)
            heading.runs[0].font.color.rgb = RGBColor(68, 114, 196)

        # Heading 3
        elif line.startswith('### '):
            text = line[4:]
            heading = doc.add_heading(text, level=3)

        # Heading 4
        elif line.startswith('#### '):
            text = line[5:]
            heading = doc.add_heading(text, level=4)

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove bold/italic markers for simplicity
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # *italic*
            text = re.sub(r'`(.*?)`', r'\1', text)        # `code`
            doc.add_paragraph(text, style='List Bullet')

        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            # Remove bold/italic markers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')

        # Code block
        elif line.startswith('```'):
            # Find end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1

            if code_lines:
                # Add code as paragraph with monospace font
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)

                # Set monospace font
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(51, 51, 51)

        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)

        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # *italic*
            text = re.sub(r'`(.*?)`', r'\1', text)        # `code`

            # Handle links [text](url) - just show text
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)

            # Skip lines that look like table separators
            if not re.match(r'^[\|\-\s]+$', text):
                doc.add_paragraph(text)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Converted {md_file} → {docx_file}")


def main():
    """Main entry point."""

    if len(sys.argv) < 2:
        print("Usage: python src/utils/md_to_docx.py <input.md> [-o output.docx]")
        print()
        print("Examples:")
        print("  python src/utils/md_to_docx.py INSTRUCTIONS.md")
        print("  python src/utils/md_to_docx.py README.md -o readme.docx")
        print()
        print("If output file not specified, uses same name with .docx extension")
        sys.exit(1)

    input_file = sys.argv[1]

    # Check if input file exists
    input_path = Path(input_file)
    if not input_path.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    # Determine output file
    if '-o' in sys.argv:
        o_index = sys.argv.index('-o')
        if o_index + 1 < len(sys.argv):
            output_file = sys.argv[o_index + 1]
        else:
            print("Error: -o flag requires output filename")
            sys.exit(1)
    else:
        # Default: replace .md with .docx
        output_file = str(input_path.with_suffix('.docx'))

    # Convert
    parse_markdown_to_docx(input_file, output_file)


if __name__ == "__main__":
    main()
